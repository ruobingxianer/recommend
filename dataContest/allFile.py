# encoding=utf-8
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfpage import PDFTextExtractionNotAllowed
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.layout import *
from pdfminer.converter import PDFPageAggregator
from docx import Document
import datetime
import os.path
import sys

reload(sys)
sys.setdefaultencoding('utf8')

path = '/Users/wanghaoxian/Desktop/数据大赛/形式审核论文'
files = os.listdir(path)
wpath = '/Users/wanghaoxian/Desktop/数据大赛/txt/'

for file in files:
    name = str(file)
    content = ""
    nowTime = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    print name, nowTime
    num = name.split("_")
    try:
        if name.endswith(".pdf"):
            fp = open(path + "/" + name, 'rb')
            parser = PDFParser(fp)
            document = PDFDocument(parser)
            if not document.is_extractable:
                raise PDFTextExtractionNotAllowed
            else:
                rsrcmgr = PDFResourceManager()
                laparams = LAParams()
                device = PDFPageAggregator(rsrcmgr, laparams=laparams)
                interpreter = PDFPageInterpreter(rsrcmgr, device)
                for page in PDFPage.create_pages(document):
                    interpreter.process_page(page)
                    layout = device.get_result()
                    for x in layout:
                        if (isinstance(x, LTTextBoxHorizontal)):
                            content = content + "\r\n" + str(x.get_text().encode('utf-8'))

        if name.endswith(".docx"):
            document = Document(path + "/" + name)
            for paragraph in document.paragraphs:
                content = content + "\r\n" + paragraph.text

    except:
        print "err"

    if(len(content)<3):
        continue

    wfile = open(wpath+num[0],"w")
    wfile.write(content)
